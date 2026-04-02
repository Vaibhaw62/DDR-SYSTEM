from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, attrs in kwargs.items():
        tag = OxmlElement(f"w:{edge}")
        for k, v in attrs.items():
            tag.set(qn(f"w:{k}"), v)
        tcBorders.append(tag)
    tcPr.append(tcBorders)


SEVERITY_COLORS = {
    "high": ("C00000", "FFFFFF"),
    "medium": ("E36C09", "FFFFFF"),
    "low": ("375623", "FFFFFF"),
}

PRIORITY_COLORS = {
    "immediate": "C00000",
    "short-term": "E36C09",
    "general maintenance": "375623",
}


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.size = Pt(10.5)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*[int(color[i:i+2], 16) for i in (0, 2, 4)])
    p.paragraph_format.space_after = Pt(3)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(str(text))
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(report, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("MAIN DETAILED DIAGNOSTIC REPORT")
    tr.bold = True
    tr.font.size = Pt(22)
    tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Prepared by AI-Powered DDR Workflow System")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    sub.paragraph_format.space_after = Pt(16)

    # Property info table
    pi = report.get("property_info", {})
    fields = [
        ("Property Address", pi.get("address", "Not Available")),
        ("Inspection Date", pi.get("inspection_date", "Not Available")),
        ("Report Date", pi.get("report_date", datetime.date.today().strftime("%d %B %Y"))),
        ("Property Type", pi.get("property_type", "Not Available")),
        ("Inspector", pi.get("inspector", "Not Available")),
        ("Client", pi.get("client", "Not Available")),
    ]
    tbl = doc.add_table(rows=len(fields), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(fields):
        row = tbl.rows[i]
        kc = row.cells[0]
        vc = row.cells[1]
        set_cell_bg(kc, "D6E4F0")
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        vp = vc.paragraphs[0]
        vr = vp.add_run(str(v))
        vr.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    add_divider(doc)

    # 1. Property Issue Summary
    add_heading(doc, "1. Property Issue Summary")
    add_body(doc, report.get("property_issue_summary", "Not Available"))
    add_divider(doc)

    # 2. Area-wise Observations
    add_heading(doc, "2. Area-wise Observations")
    for area in report.get("area_wise_observations", []):
        add_heading(doc, area.get("area", "Area"), level=2)
        for obs in area.get("observations", []):
            add_bullet(doc, obs)
        thermal = area.get("thermal_findings", [])
        if thermal:
            p = doc.add_paragraph()
            r = p.add_run("Thermal Findings:")
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x88, 0x44, 0x00)
            for t in thermal:
                add_bullet(doc, t)
        for img in area.get("images", []):
            na = not img or "not available" in str(img).lower()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("[ Image Not Available ]" if na else f"[ Image: {img} ]")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
    add_divider(doc)

    # 3. Probable Root Causes
    add_heading(doc, "3. Probable Root Causes")
    for cause in report.get("probable_root_causes", ["Not Available"]):
        add_bullet(doc, cause)
    add_divider(doc)

    # 4. Severity Assessment
    add_heading(doc, "4. Severity Assessment")
    sev_data = report.get("severity_assessment", [])
    if sev_data:
        tbl2 = doc.add_table(rows=len(sev_data) + 1, cols=3)
        tbl2.style = "Table Grid"
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["Area / Issue", "Severity", "Reasoning"]
        hrow = tbl2.rows[0]
        for j, h in enumerate(headers):
            cell = hrow.cells[j]
            set_cell_bg(cell, "1F4E79")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, s in enumerate(sev_data):
            row = tbl2.rows[i + 1]
            lvl = s.get("severity", "Low").lower()
            bg, fg = SEVERITY_COLORS.get(lvl, ("375623", "FFFFFF"))

            area_cell = row.cells[0]
            sev_cell = row.cells[1]
            reas_cell = row.cells[2]

            p1 = area_cell.paragraphs[0]
            r1 = p1.add_run(f"{s.get('issue','')} ({s.get('area','')})")
            r1.font.size = Pt(10)

            set_cell_bg(sev_cell, bg)
            p2 = sev_cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(s.get("severity", "Low"))
            r2.bold = True
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(*[int(fg[i:i+2], 16) for i in (0, 2, 4)])

            p3 = reas_cell.paragraphs[0]
            r3 = p3.add_run(s.get("reasoning", ""))
            r3.font.size = Pt(10)
    add_divider(doc)

    # 5. Recommended Actions
    add_heading(doc, "5. Recommended Actions")
    groups = {}
    for a in report.get("recommended_actions", []):
        grp = a.get("priority", "General Maintenance")
        groups.setdefault(grp, []).append(a.get("action", ""))
    priority_order = ["Immediate", "Short-term", "General Maintenance"]
    for grp in priority_order:
        items = groups.get(grp, [])
        if not items:
            continue
        col = PRIORITY_COLORS.get(grp.lower(), "333333")
        p = doc.add_paragraph()
        r = p.add_run(grp.upper())
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(*[int(col[i:i+2], 16) for i in (0, 2, 4)])
        for i, item in enumerate(items, 1):
            p2 = doc.add_paragraph(style="List Number")
            rr = p2.add_run(item)
            rr.font.size = Pt(10)
    add_divider(doc)

    # 6. Additional Notes
    add_heading(doc, "6. Additional Notes")
    for note in report.get("additional_notes", ["Not Available"]):
        add_bullet(doc, note)
    conflicts = report.get("conflicts", [])
    if conflicts:
        p = doc.add_paragraph()
        r = p.add_run("Conflicts Between Documents:")
        r.bold = True
        r.font.color.rgb = RGBColor(0xE3, 0x6C, 0x09)
        r.font.size = Pt(10)
        for c in conflicts:
            add_bullet(doc, c)
    add_divider(doc)

    # 7. Missing or Unclear Information
    add_heading(doc, "7. Missing or Unclear Information")
    missing = report.get("missing_or_unclear_information", [])
    if missing:
        tbl3 = doc.add_table(rows=len(missing) + 1, cols=2)
        tbl3.style = "Table Grid"
        tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER

        for j, h in enumerate(["Item", "Status"]):
            cell = tbl3.rows[0].cells[j]
            set_cell_bg(cell, "1F4E79")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for i, m in enumerate(missing):
            row = tbl3.rows[i + 1]
            p1 = row.cells[0].paragraphs[0]
            p1.add_run(m.get("item", "")).font.size = Pt(10)
            p2 = row.cells[1].paragraphs[0]
            r2 = p2.add_run(m.get("status", "Not Available"))
            r2.italic = True
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    else:
        add_body(doc, "None identified.", italic=True)

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— END OF REPORT —")
    er.bold = True
    er.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    er.font.size = Pt(10)

    doc.save(output_path)
